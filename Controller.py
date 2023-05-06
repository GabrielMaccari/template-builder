# -*- coding: utf-8 -*-
"""
@author: Gabriel Maccari
"""

import pandas
import docx

COLUNAS_TABELA_CADERNETA = {
    "Ponto": {
        "dtype": "object", "nulo_ok": False, "dominio": None
    },
    "Disciplina": {
        "dtype": "object", "nulo_ok": False,
        "dominio": ["Mapeamento Geológico I", "Mapeamento Geológico II"]
    },
    "SRC": {
        "dtype": "object", "nulo_ok": False, "dominio": None
    },
    "Easting": {
        "dtype": "float64", "nulo_ok": False, "dominio": None
    },
    "Northing": {
        "dtype": "float64", "nulo_ok": False, "dominio": None
    },
    "Altitude": {
        "dtype": "float64", "nulo_ok": True, "dominio": None
    },
    "Toponimia": {
        "dtype": "object", "nulo_ok": True, "dominio": None
    },
    "Data": {
        "dtype": "datetime64[ns]", "nulo_ok": False, "dominio": None
    },
    "Equipe": {
        "dtype": "object", "nulo_ok": False, "dominio": None
    },
    "Ponto_de_controle": {
        "dtype": "object", "nulo_ok": False, "dominio": ["Sim", "Não"]
    },
    "Numero_de_amostras": {
        "dtype": "int64", "nulo_ok": False, "dominio": None
    },
    "Possui_croquis": {
        "dtype": "object", "nulo_ok": False, "dominio": ["Sim", "Não"]
    },
    "Possui_fotos": {
        "dtype": "object", "nulo_ok": False, "dominio": ["Sim", "Não"]
    },
    "Tipo_de_afloramento": {
        "dtype": "object", "nulo_ok": True, "dominio": None
    },
    "In_situ": {
        "dtype": "object", "nulo_ok": True, "dominio": ["Sim", "Não"]
    },
    "Grau_de_intemperismo": {
        "dtype": "object", "nulo_ok": True, "dominio": ["Baixo", "Médio", "Alto"]
    },
    "Unidade": {
        "dtype": "object", "nulo_ok": True, "dominio": None
    },
    "Unidade_litoestratigrafica": {
        "dtype": "object", "nulo_ok": True, "dominio": None
    }
}


class ControladorPrincipal:
    def __init__(self, caminho_template: str, df: pandas.DataFrame = None):
        self.caminho_template = caminho_template
        self.template = docx.Document(caminho_template)
        self.df = df
        self.caderneta = None

        self.estilos = {
            "normal": self.template.styles['Normal'],
            "titulo": self.template.styles['Title'],
            "titulo1": self.template.styles['Heading 1'],
            "titulo2": self.template.styles['Heading 2'],
            "subtitulo": self.template.styles['Subtitle'],
            "titulo_informacao": self.template.styles['Título de informação'],
            "texto_informacao": self.template.styles['Texto de informação'],
            "legenda": self.template.styles['Caption'],
            "tabela_esquerda": self.template.styles['Tabela - Coluna esquerda'],
            "tabela_direita": self.template.styles['Tabela - Coluna direita'],
            "tabela_cabecalho": self.template.styles['Tabela de cabeçalho'],
        }

    def recarregar_template(self):
        self.template = None
        self.template = docx.Document(self.caminho_template)

    def abrir_tabela(self, caminho: str) -> object:
        """
        Abre uma tabela do excel e armazena o DataFrame no atributo "df" do controlador.
        :param caminho: O caminho até um arquivo .xlsx ou .xlsm.
        :returns: Boolean dizendo se o DataFrame foi criado com sucesso e Integer com o número de linhas do DataFrame
        """
        # Salva a primeira aba da tabela em um DataFrame
        df = pandas.read_excel(caminho, engine='openpyxl')
        # Descarta colunas sem nome
        colunas_remocao = [col for col in df.columns if 'Unnamed' in col]
        df.drop(colunas_remocao, axis='columns', inplace=True)
        # Descarta linhas vazias
        df.dropna(how='all', axis='index', inplace=True)
        # Verifica se existem linhas preenchidas no arquivo
        linhas = len(df.index)
        if linhas <= 0:
            raise Exception('A tabela selecionada está vazia ou contém apenas cabeçalhos.')
        # Checa se o dataframe foi criado ou não e armazena no model
        if isinstance(df, pandas.DataFrame):
            self.df = df
            self.caderneta = None
            return True, linhas
        else:
            return False, linhas

    def checar_colunas(self) -> list[str]:
        """
        Checa se cada coluna esperada para a tabela existe, está no formato correto, contém apenas valores permitidos.
        O DataFrame é obtido do atributo "df" do controlador.
        :returns: Lista de strings especificando o status de cada coluna. O status pode ser "ok", "faltando", "problemas", "nulos" ou "dominio"
        """
        df = self.df
        colunas_df = df.columns.to_list()

        status_colunas = []
        for c in COLUNAS_TABELA_CADERNETA:
            dtype = COLUNAS_TABELA_CADERNETA[c]["dtype"]
            nulo_ok = COLUNAS_TABELA_CADERNETA[c]["nulo_ok"]
            dominio = COLUNAS_TABELA_CADERNETA[c]["dominio"]

            # Checa se a coluna existe na tabela
            if c not in colunas_df:
                status_colunas.append("missing_column")
                continue

            # Verifica se existem nulos e se a coluna permite nulos
            if not nulo_ok and df[c].isnull().values.any():
                status_colunas.append("nan_not_allowed")
                continue

            # Tenta converter a tabela para o tipo de dado esperado
            try:
                df[c] = df[c].astype(dtype, errors="raise")
            except ValueError:
                status_colunas.append("wrong_dtype")
                continue

            # Verifica se a coluna possui valores controlados e se existe algum
            # valor fora do domínio
            if dominio is not None:
                valores_coluna = df[c]
                if nulo_ok:
                    valores_coluna.dropna(inplace=True)
                if not valores_coluna.isin(dominio).all():
                    status_colunas.append("outside_domain")
                    continue

            status_colunas.append("ok")

        return status_colunas

    def localizar_problemas_formato(self, coluna: str) -> list[int]:
        """
        Localiza as linhas da tabela com problemas que impedem a conversão para o tipo de dado esperado.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os indexes das linhas com problema.
        """
        valores = self.df[coluna].dropna()
        tipo_alvo = COLUNAS_TABELA_CADERNETA[coluna]["dtype"]

        funcoes_conversao = {
            "datetime64[ns]": pandas.to_datetime(valores, errors="coerce", format="%d/%m/%Y").isna(),
            "float64": pandas.to_numeric(valores, errors="coerce", downcast="float").isna(),
            "int64": pandas.to_numeric(valores, errors="coerce", downcast="integer").isna()
        }

        if tipo_alvo not in funcoes_conversao:
            raise Exception(f"Checagem não implementada para o tipo de dado ({tipo_alvo})")

        # Valores que não podem ser convertidos tornam-se NaN devido ao "coerce"
        convertido = funcoes_conversao[tipo_alvo]
        indices_problemas = [i for i, is_nan in zip(convertido.index, convertido.values) if is_nan]
        return indices_problemas

    def localizar_celulas_vazias(self, coluna: str) -> list[int]:
        """
        Localiza as linhas da coluna especificada que contêm valores nulos.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os indexes das linhas com problema.
        """
        valores_coluna = self.df.loc[:, coluna]
        indices_problemas = self.df[valores_coluna.isnull()].index.tolist()
        return indices_problemas

    def localizar_problemas_dominio(self, coluna: str) -> list[int]:
        """
        Localiza células em uma coluna com valores fora de domínio.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os indexes das linhas com problema.
        """
        valores_coluna = self.df.loc[:, coluna]
        dominio = COLUNAS_TABELA_CADERNETA[coluna]["dominio"]
        indices_problemas = valores_coluna.index[~valores_coluna.isin(dominio)].tolist()
        return indices_problemas

    def montar_msg_problemas(self, tipo_problema: str, coluna: str, indices: list[int]) -> str:
        """
        Monta a mensagem especificando quais linhas da tabela estão com problemas.
        :param tipo_problema: "missing_column", "wrong_dtype", "nan_not_allowed" ou "outside_domain"
        :param coluna: O nome da coluna.
        :param indices: Os índices das linhas com problemas no DataFrame.
        :returns: String descrevendo o problema e as linhas que devem ser corrigidas.
        """
        dtype_coluna = str(COLUNAS_TABELA_CADERNETA[coluna]["dtype"])

        tipos_problemas = {
            "missing_column": (
                f"A coluna \"{coluna}\" não foi encontrada na tabela. "
                f"Verifique se ela foi excluída ou se você selecionou a tabela errada. "
                f"Restaure a coluna ou tente novamente com a tabela correta."
            ),
            "wrong_dtype": (
                f"A coluna \"{coluna}\" possui dados fora do formato aceito ({dtype_coluna}) "
                f"nas linhas especificadas abaixo. Corrija-os e tente novamente.\n"
            ),
            "nan_not_allowed": (
                f"Existem células vazias nas seguintes linhas da coluna \"{coluna}\". "
                f"Preencha apropriadamente as células em questão e tente novamente.\n"
            ),
            "outside_domain": (
                f"A coluna \"{coluna}\" possui valores fora da lista de valores permitidos "
                f"nas seguintes linhas. Corrija-os e tente novamente.\n"
            )
        }

        mensagem = [tipos_problemas.get(tipo_problema)]

        for i in indices:
            linha = i + 2
            ponto = self.df.loc[i, ["Ponto"]].values[0]
            mensagem.append(f"Linha {linha} (ponto {ponto})")

        return "\n".join(mensagem)

    def gerar_caderneta(self, montar_folha_de_rosto: bool = True):
        """
        Gera a caderneta pré-preenchida.
        :param folha_de_rosto: Opção para gerar ou não uma folha de rosto.
        :returns: Nada.
        """
        # Limpa todos os objetos da classe docx.Document para evitar bugs comuns
        self.recarregar_template()
        self.caderneta = None
        documento = None

        documento = self.template
        df = self.df
        colunas_tabela = df.columns.to_list()

        # Na tabela da caderneta, as colunas 19-33 são potenciais colunas de medidas estruturais
        colunas_estrutura = (colunas_tabela[18:] if len(colunas_tabela) < 33
                             else colunas_tabela[18:33])

        try:
            df['Data'] = df['Data'].dt.strftime('%d/%m/%Y')
        except:
            pass

        df["Possui_croquis"] = df["Possui_croquis"].map({"Sim": True, "Não": False})
        df["Possui_fotos"] = df["Possui_fotos"].map({"Sim": True, "Não": False})

        # Deleta o primeiro parágrafo do template (aviso para não excluir o arquivo)
        paragraph = documento.paragraphs[0]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

        # Monta a folha de rosto da caderneta
        if montar_folha_de_rosto:
            documento = self.montar_folha_rosto(documento)

        d = 1  # Número sequencial do mestre/disciplina. Ex: Map1 = 1
        disciplinas = COLUNAS_TABELA_CADERNETA["Disciplina"]["dominio"]

        for linha in df.itertuples():
            # Adiciona uma página de título antes do primeiro ponto de cada semestre/disciplina
            if d <= 2 and linha.Disciplina == disciplinas[d-1]:
                documento = self.montar_pagina_semestre(documento, linha.Disciplina)
                d += 1

            # Quebra a página antes do título do ponto
            documento.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

            # Adiciona a página do ponto
            documento = self.montar_pagina_ponto(documento, linha, colunas_estrutura)

        self.caderneta = documento

    def montar_folha_rosto(self, documento: docx.Document) -> docx.Document:
        """
        Adiciona uma folha de rosto à caderneta.
        :param documento: O documento.
        :returns: O documento com a folha de rosto.
        """
        for i in range(0, 15):
            if i == 10:
                documento.add_paragraph(text='CADERNETA DE CAMPO COMPILADA',
                                        style=self.estilos["titulo"])
            elif i == 13:
                documento.add_paragraph(text='MAPEAMENTO GEOLÓGICO UFSC',
                                        style=self.estilos["titulo_informacao"])
            else:
                documento.add_paragraph(text='', style=self.estilos['normal'])

        lista_infos = ['PROJETO:', 'ANO:', 'PROFESSORES RESPONSÁVEIS:',
                       'NÚMERO DA ÁREA/FAIXA:', 'INTEGRANTES DO GRUPO:']

        for info in lista_infos:
            documento.add_paragraph(text=info, style=self.estilos["titulo_informacao"])
            documento.add_paragraph(text='<PREENCHA AQUI>', style=self.estilos["texto_informacao"])

        return documento

    def montar_pagina_semestre(self, documento: docx.Document, disciplina: str) -> docx.Document:
        """
        Adiciona uma página de título à caderneta para dividir os semestres do mapeamento geológico.
        :param documento: O documento.
        :param disciplina: "Mapeamento Geológico I" ou "Mapeamento Geológico II".
        :returns: O documento com a página de título do semestre.
        """
        try:  # Quando não há folha de rosto, o documento está inicialmente vazio, e isso causa um IndexError
            documento.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        except IndexError:
            pass
        for i in range(0, 18):
            documento.add_paragraph(text='', style=self.estilos["normal"])
        documento.add_heading(text=disciplina, level=1)

        return documento

    def montar_pagina_ponto(self, documento: docx.Document, linha: pandas.core.frame.pandas,
                            colunas_estrutura: list[str]) -> docx.Document:
        """
        Acrescenta uma página de informações de um ponto à caderneta.
        :param documento: O documento
        :param linha: Duplas de rótulos e valores da linha do DataFrame (gerado via DataFrame.itertuples().
        :param colunas_estrutura: Os nomes das colunas de medidas estruturais presentes na tabela.
        :returns: O documento com a página do ponto.
        """

        # Valores das colunas para a linha
        ponto = linha.Ponto
        src = linha.SRC
        easting = linha.Easting
        northing = linha.Northing
        altitude = linha.Altitude
        toponimia = linha.Toponimia
        data = linha.Data
        equipe = linha.Equipe
        ponto_controle = linha.Ponto_de_controle
        num_amostras = linha.Numero_de_amostras
        possui_croquis = linha.Possui_croquis
        possui_fotos = linha.Possui_fotos
        tipo_afloramento = linha.Tipo_de_afloramento
        in_situ = linha.In_situ
        intemperismo = linha.Grau_de_intemperismo
        unidade = linha.Unidade
        unidade_lito = linha.Unidade_litoestratigrafica

        # Título do ponto
        documento.add_heading(text=ponto, level=2)

        # Dicionário com informações que irão para a tabela de cabeçalho
        dados_tabela = {
            'DATA:': f"{data}",
            'COORDENADAS:': f"{easting:.0f}E {northing:.0f}N   {src}",
            'ALTITUDE:': f"{altitude:.0f} m" if not pandas.isna(altitude) else "-",
            'TOPONÍMIA:': f"{toponimia}" if not pandas.isna(toponimia) else "-",
            'EQUIPE:': f"{equipe}",
            'PONTO DE CONTROLE:': f"{ponto_controle}",
            'TIPO DE AFLORAMENTO:': f"{tipo_afloramento}" if not pandas.isna(tipo_afloramento) else "-",
            'IN SITU:': f"{in_situ}" if not pandas.isna(in_situ) else "-",
            'GRAU DE INTEMPERISMO:': f"{intemperismo}" if not pandas.isna(intemperismo) else "-",
            'AMOSTRAS:': f"{num_amostras}" if num_amostras > 0 else "-",
            'UNIDADE:': f"{unidade} - {unidade_lito}" if not pandas.isna(unidade) else "-"
        }

        # Preenche a tabela de cabeçalho
        table = documento.add_table(rows=0, cols=2)
        table.style = self.estilos["tabela_cabecalho"]
        for key in dados_tabela.keys():
            lin = table.add_row().cells
            # Coluna esquerda
            lin[0].text = key
            lin[0].paragraphs[0].style = self.estilos["tabela_esquerda"]
            # Coluna direita
            lin[1].text = str(dados_tabela[key]) if str(dados_tabela[key]) != 'nan' else '-'
            lin[1].paragraphs[0].style = self.estilos["tabela_direita"]

        # Ajusta a largura das colunas da tabela
        for celula in table.columns[0].cells:
            celula.width = docx.shared.Inches(2.1)
        for celula in table.columns[1].cells:
            celula.width = docx.shared.Inches(3.8)

        # Adiciona a seção de descrição do ponto
        documento.add_paragraph(text='DESCRIÇÃO', style=self.estilos["subtitulo"])
        documento.add_paragraph(text="<Descrição do afloramento aqui>", style=self.estilos["normal"])

        # Se for um ponto de controle, encerra aqui
        if ponto_controle == "Sim":
            return documento

        # Adiciona a seção de amostras, se houver alguma
        if num_amostras > 0:
            documento.add_paragraph(text='AMOSTRAS', style=self.estilos["subtitulo"])
            abc = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
            for i in range(0, num_amostras):
                letra = abc[i]
                documento.add_paragraph(text=f"• {ponto}{letra}: <Descrição da amostra aqui>", style=self.estilos["normal"])

        # Procura medidas estruturais na tabela
        medidas_estruturais = []
        for i, coluna in enumerate(colunas_estrutura):
            # Conteúdo do campo
            medida = linha[i + 19]
            # Se não for uma célula vazia
            if not pandas.isna(medida):
                # Procura uma sigla entre parênteses
                if '(' in coluna and ')' in coluna:
                    sigla = coluna[coluna.find("(") + 1:coluna.find(")")]
                # Se não encontrar sigla, usa o nome da coluna
                else:
                    sigla = coluna.replace('_', ' ')
                # Adiciona as medidas a uma lista
                medidas_estruturais.append(f"• {sigla} = {medida}")

        # Adiciona a seção de medidas, se houver alguma
        if len(medidas_estruturais) > 0:
            documento.add_paragraph(text='MEDIDAS ESTRUTURAIS', style=self.estilos["subtitulo"])
            for m in medidas_estruturais:
                documento.add_paragraph(text=m, style=self.estilos["normal"])

        # Adiciona a seção de croquis, se houver algum
        if possui_croquis:
            documento.add_paragraph(text='CROQUIS', style=self.estilos["subtitulo"])
            documento.add_paragraph(
                text="<Insira aqui os croquis elaborados para o afloramento e suas "
                     "respectivas legendas. Remova esta seção caso não haja croquis>",
                style=self.estilos["normal"]
            )

        # Adiciona a seção de fotos, se houver alguma
        if possui_fotos:
            documento.add_paragraph(text='FOTOS', style=self.estilos["subtitulo"])
            documento.add_paragraph(
                text="<Insira aqui os painéis de fotos tiradas no afloramento e suas "
                     "respectivas legendas. Remova esta seção caso não haja fotos>",
                style=self.estilos["normal"]
            )

        return documento

    def salvar_caderneta(self, caminho: str):
        """
        Salva a caderneta como um arquivo .docx.
        :param caminho: O caminho do arquivo.
        :returns: Nada.
        """
        if not caminho.endswith(".docx"):
            caminho += ".docx"
        self.caderneta.save(caminho)
