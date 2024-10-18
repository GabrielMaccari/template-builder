# -*- coding: utf-8 -*-
""" @author: Gabriel Maccari """

import json
import pandas
import docx
import numpy
from datetime import datetime
from icecream import ic

from Controller import mostrar_popup

COLUNAS_ABA_LISTAS = (
    "Unidades geológicas",
    "Unidades litoestratigráficas",
    "Estruturas planares",
    "Estruturas lineares",
    "Áreas/Faixas",
    "Fases"
)


class Modelo:
    def __init__(self, caminho_template: str, caminho_colunas: str, df: pandas.DataFrame = None):
        self.df = df
        self.caminho_template = None
        self.template = None
        self.colunas = None
        self.estilos = None
        self.caderneta = None

        self.carregar_template(caminho_template)
        self.carregar_colunas(caminho_colunas)

    def carregar_template(self, caminho: str):
        """
        Carrega o template para a caderneta a partir de um arquivo .docx.
        :param caminho: Caminho para o documento a ser usado como template.
        :returns: Nada.
        """
        ic(caminho)

        self.caminho_template = caminho

        # Essa linha evitava um bug que stackava cadernetas quando mais de uma era gerada na mesma execução da
        # ferramenta. Eu não sei se esse bug ainda acontece, mas deixei a linha aqui por garantia
        self.template = None

        self.template = docx.Document(caminho)
        self.estilos = {
            "normal": self.template.styles['Normal'],
            "titulo": self.template.styles['Title'],
            "titulo1": self.template.styles['Heading 1'],
            "titulo2": self.template.styles['Heading 2'],
            "subtitulo": self.template.styles['Subtitle'],
            "titulo_informacao": self.template.styles['Título de informação'],
            "texto_informacao": self.template.styles['Texto de informação'],
            "anotacao": self.template.styles['Anotação'],
            "legenda": self.template.styles['Caption'],
            "tabela_esquerda": self.template.styles['Tabela - Coluna esquerda'],
            "tabela_direita": self.template.styles['Tabela - Coluna direita'],
            "tabela_cabecalho": self.template.styles['Tabela de cabeçalho'],
        }

    def carregar_colunas(self, caminho: str):
        """
        carrega o arquivo JSON de definições das colunas da aba Geral.
        :param caminho: Caminho para o arquivo JSON contendo as definições das colunas.
        :returns: Nada.
        """
        ic(caminho)

        with open(caminho, 'r', encoding='utf-8') as arquivo:
            self.colunas = json.load(arquivo)

    def abrir_tabela(self, caminho: str) -> object:
        """
        Abre a planilha de pontos e lê a aba Geral para extrair as informações para a caderneta.
        Armazena o DataFrame no atributo "df" do controlador.
        :param caminho: O caminho até um arquivo .xlsx ou .xlsm.
        :returns: Boolean dizendo se o DataFrame foi criado com sucesso e Integer com o número de linhas do DataFrame
        """
        ic(caminho)

        # Salva a primeira aba da tabela em um DataFrame
        df = pandas.read_excel(caminho, sheet_name="Geral", engine='openpyxl')
        # Converte os nomes das colunas para string
        df.columns = df.columns.astype(str)
        # Descarta colunas sem nome
        colunas_remocao = [col for col in df.columns if 'Unnamed' in col]
        df.drop(colunas_remocao, axis='columns', inplace=True)
        # Descarta linhas vazias
        df.dropna(how='all', axis='index', inplace=True)

        ic(df.dtypes)

        # Verifica se existem linhas preenchidas no arquivo
        linhas = len(df.index)
        if linhas <= 0:
            raise Exception('A tabela selecionada está vazia ou contém apenas cabeçalhos.')
        # Checa se o dataframe foi criado ou não e armazena no model
        if isinstance(df, pandas.DataFrame):
            self.df = df
            self.caderneta = None
            return True, linhas
        else:
            return False, linhas

    def abrir_aba_listas(self, caminho):
        """
        Lê a aba Listas da tabela de pontos e usa as listas de opções para especificar o domínio das colunas.
        :param caminho: O caminho até um arquivo .xlsx ou .xlsm.
        :returns: Nada.
        """
        ic(caminho)
        df_listas = pandas.read_excel(caminho, sheet_name="Listas", engine='openpyxl')

        try:
            lista_und_geo = df_listas["Unidades geológicas"].head(30).dropna().to_list()
            lista_und_lito = df_listas["Unidades litoestratigráficas"].head(30).dropna().to_list()
            lista_faixas = df_listas["Áreas/Faixas"].head(30).dropna().to_list()
            lista_fases = df_listas["Fases"].head(30).dropna().to_list()
        except Exception as e:
            ic(e)
            colunas_problema = [c for c in COLUNAS_ABA_LISTAS if c not in df_listas.columns]
            raise Exception(f"Uma ou mais colunas essenciais estão faltando ou tiveram o cabeçalho modificado na aba de"
                            f" Listas da tabela:\n\n{", ".join(colunas_problema)}")

        self.colunas["Unidade_geologica_1"]["dominio"] = lista_und_geo
        self.colunas["Unidade_geologica_2"]["dominio"] = lista_und_geo
        self.colunas["Unidade_litoestratigrafica_1"]["dominio"] = lista_und_lito
        self.colunas["Unidade_litoestratigrafica_2"]["dominio"] = lista_und_lito
        self.colunas["Faixa"]["dominio"] = lista_faixas
        self.colunas["Fase"]["dominio"] = lista_fases

    def checar_colunas(self) -> list[str]:
        """
        Checa se cada coluna esperada para a tabela existe, está no formato correto, contém apenas valores permitidos.
        O DataFrame é obtido do atributo "df" do controlador.
        :returns: Lista de strings especificando o status de cada coluna. O status pode ser "ok", "faltando", "problemas", "nulos" ou "dominio"
        """
        ic()

        df = self.df
        colunas_df = df.columns.to_list()

        status_colunas = []
        for c in self.colunas:
            dtype = self.colunas[c]["dtype"]
            nulo_ok = self.colunas[c]["nulo_ok"]
            dominio = self.colunas[c]["dominio"]
            intervalo = self.colunas[c]["intervalo"]
            unico = self.colunas[c]["unico"]

            # Checa se a coluna existe na tabela
            if c not in colunas_df:
                status_colunas.append("coluna_faltando")
                continue

            # Verifica se existem nulos e se a coluna permite nulos
            if not nulo_ok and df[c].isnull().values.any():
                status_colunas.append("celulas_vazias")
                continue

            # Tenta converter a tabela para o tipo de dado esperado
            try:
                df[c] = df[c].astype(dtype, errors="raise")
            except ValueError:
                status_colunas.append("fora_de_formato")
                continue

            # Verifica se a coluna possui valores controlados e se existe algum valor fora do domínio
            if dominio is not None:
                valores_coluna = df[c]
                if nulo_ok:
                    valores_coluna.dropna(inplace=True)
                if not valores_coluna.isin(dominio).all():
                    status_colunas.append("valores_nao_permitidos")
                    continue

            # Verifica se valores numéricos da coluna estão dentro do intervalo esperado
            if intervalo is not None:
                valores_coluna = df[c]
                if nulo_ok:
                    valores_coluna.dropna(inplace=True)
                if not valores_coluna.between(intervalo[0], intervalo[1]).all():
                    status_colunas.append("fora_do_intervalo")
                    continue

            # Checa se existem valores repetidos não-permitidos na coluna
            if unico and not df[c].nunique() == df[c].count():
                status_colunas.append("valores_repetidos")
                continue

            status_colunas.append("ok")

        return status_colunas

    def localizar_valores_repetidos(self, coluna: str) -> list[int]:
        """
        Localiza as linhas da coluna especificada onde há valores repetidos/não-únicos.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os índices das linhas com problema.
        """
        ic(coluna)

        duplicados = self.df[self.df[coluna].duplicated(keep=False)]
        duplicados = duplicados[coluna]
        indices_problemas = [i for i, duplicado in zip(duplicados.index, duplicados.values) if duplicado]
        return indices_problemas

    def localizar_problemas_formato(self, coluna: str) -> list[int]:
        """
        Localiza as linhas da tabela com problemas que impedem a conversão para o tipo de dado esperado.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os índices das linhas com problema.
        """
        ic(coluna)

        valores_coluna = self.df[coluna].dropna()
        tipo_alvo = self.colunas[coluna]["dtype"]

        funcoes_conversao = {
            "datetime64[ns]": pandas.to_datetime(valores_coluna, errors="coerce", format="%d/%m/%Y").isna(),
            "float64": pandas.to_numeric(valores_coluna, errors="coerce", downcast="float").isna(),
            "int64": pandas.to_numeric(valores_coluna, errors="coerce", downcast="integer").isna()
        }

        if tipo_alvo not in funcoes_conversao:
            raise Exception(f"Checagem não implementada para o tipo de dado presente na coluna {coluna} ({tipo_alvo}).")

        # Valores que não podem ser convertidos tornam-se NaN devido ao "coerce"
        convertido = funcoes_conversao[tipo_alvo]
        indices_problemas = [i for i, is_nan in zip(convertido.index, convertido.values) if is_nan]
        return indices_problemas

    def localizar_celulas_vazias(self, coluna: str) -> list[int]:
        """
        Localiza as linhas da coluna especificada que contêm valores nulos.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os índices das linhas com problema.
        """
        ic(coluna)

        valores_coluna = self.df.loc[:, coluna]
        indices_problemas = self.df[valores_coluna.isnull()].index.tolist()
        return indices_problemas

    def localizar_problemas_dominio(self, coluna: str) -> list[int]:
        """
        Localiza células em uma coluna com valores fora de domínio.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os índices das linhas com problema.
        """
        ic(coluna)

        valores_coluna = self.df.loc[:, coluna]
        dominio = self.colunas[coluna]["dominio"]
        indices_problemas = valores_coluna.index[~valores_coluna.isin(dominio)].tolist()
        return indices_problemas

    def localizar_problemas_intervalo(self, coluna: str) -> list[int]:
        """
        Localiza células em uma coluna com valores numéricos fora do intervalo permitido.
        :param coluna: O nome da coluna a ser verificada.
        :returns: Lista contendo os índices das linhas com problema.
        """
        ic(coluna)

        valores_coluna = self.df.loc[:, coluna]
        intervalo = self.colunas[coluna]["intervalo"]
        indices_problemas = valores_coluna.index[~valores_coluna.between(intervalo[0], intervalo[1])].tolist()
        return indices_problemas

    def montar_msg_problemas(self, tipo_problema: str, coluna: str, indices: list[int]) -> str:
        """
        Monta a mensagem especificando quais linhas da tabela estão com problemas.
        :param tipo_problema: "missing_column", "wrong_dtype", "nan_not_allowed" ou "outside_domain"
        :param coluna: O nome da coluna.
        :param indices: Os índices das linhas com problemas no DataFrame.
        :returns: String descrevendo o problema e as linhas que devem ser corrigidas.
        """
        ic(tipo_problema, coluna, indices)

        dtype_coluna = str(self.colunas[coluna]["dtype"])

        tipos_problemas = {
            "coluna_faltando": (
                f"A coluna \"{coluna}\" não foi encontrada na tabela. "
                f"Verifique se ela foi excluída ou se você selecionou a tabela errada. "
                f"Restaure a coluna ou tente novamente com a tabela correta."
            ),
            "fora_de_formato": (
                f"A coluna \"{coluna}\" possui dados fora do formato aceito ({dtype_coluna}) "
                f"ou fora dos limites esperados para o tipo de dado "
                f"nas linhas especificadas abaixo. Corrija-os e tente novamente.\n"
            ),
            "celulas_vazias": (
                f"Existem células vazias nas seguintes linhas da coluna \"{coluna}\". "
                f"Preencha apropriadamente as células em questão e tente novamente.\n"
            ),
            "valores_nao_permitidos": (
                f"A coluna \"{coluna}\" possui valores fora da lista de valores permitidos "
                f"nas seguintes linhas. Corrija-os e tente novamente.\n"
            ),
            "fora_do_intervalo": (
                f"A coluna \"{coluna}\" possui valores fora do intervalo numérico permitido "
                f"nas seguintes linhas. Corrija-os e tente novamente.\n"
            ),
            "valores_repetidos": (
                f"A coluna \"{coluna}\" possui valores repetidos nas seguintes linhas. "
                f"Corrija-os e tente novamente.\n"
            )
        }

        mensagem = [tipos_problemas.get(tipo_problema)]

        for i in indices:
            linha = i + 2
            ponto = self.df.loc[i, ["Ponto"]].values[0]
            mensagem.append(f"Linha {linha} (ponto {ponto})")

        return "\n".join(mensagem)

    def gerar_caderneta(self, montar_folha_de_rosto: bool = True, montar_folhas_fase: bool = True,
                        indice_inicio: int | pandas.Index = 0, continuar_caderneta: str = None):
        """
        Gera a caderneta pré-preenchida.
        :param montar_folha_de_rosto: Opção para gerar ou não uma folha de rosto.
        :param montar_folhas_fase: Opção para gerar ou não páginas de título das disciplinas.
        :param indice_inicio: O índice do DataFrame (ponto) no qual a montagem da caderneta deve iniciar.
        :param continuar_caderneta: O caminho para uma caderneta pré-existente a ser continuada (.docx). Opcional.
        :returns: Nada.
        """
        ic(montar_folha_de_rosto, montar_folhas_fase, indice_inicio, continuar_caderneta)

        # Essa linha evitava um bug que stackava cadernetas quando mais de uma era gerada na mesma execução da
        # ferramenta. Eu não sei se esse bug ainda acontece, mas deixei a linha aqui por garantia
        documento = None

        self.caderneta = None

        self.carregar_template(self.caminho_template if not continuar_caderneta else continuar_caderneta)
        documento = self.template

        if not continuar_caderneta:
            # Deleta o primeiro parágrafo do template (aquele aviso para não excluir o arquivo)
            paragraph = documento.paragraphs[0]
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

        df = self.df
        colunas_tabela = df.columns.to_list()

        # Formata as datas
        try:
            df['Data'] = df['Data'].dt.strftime('%d/%m/%Y')
        except AttributeError:
            pass

        fase = None

        # Monta a folha de rosto da caderneta
        if montar_folha_de_rosto:
            documento = self.montar_folha_rosto(documento)

        for linha in df.itertuples():
            # Pula linhas até chegar ao ponto de início
            if linha.Index < indice_inicio:
                continue

            if montar_folhas_fase:
                # Adiciona uma página de título antes do primeiro ponto de cada semestre/disciplina
                if linha.Fase != fase:
                    fase = linha.Fase
                    self.montar_pagina_fase(documento, fase)

            # Quebra a página antes do título do ponto
            documento.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

            # Adiciona a página do ponto
            documento = self.montar_pagina_ponto(documento, linha)

        self.caderneta = documento

    def montar_folha_rosto(self, documento: docx.Document) -> docx.Document:
        """
        Adiciona uma folha de rosto à caderneta.
        :param documento: O documento.
        :returns: O documento com a folha de rosto.
        """
        ic()

        for i in range(0, 15):
            if i == 10:
                documento.add_paragraph(text='CADERNETA DE CAMPO COMPILADA',
                                        style=self.estilos["titulo"])
            elif i == 13:
                documento.add_paragraph(text='MAPEAMENTO GEOLÓGICO',
                                        style=self.estilos["titulo_informacao"])
            else:
                documento.add_paragraph(text='', style=self.estilos['normal'])

        lista_infos = ['PROJETO:', 'ANO:', 'PROFESSORES RESPONSÁVEIS:',
                       'ÁREA/FAIXA:', 'INTEGRANTES DO GRUPO:']

        for info in lista_infos:
            documento.add_paragraph(text=info, style=self.estilos["titulo_informacao"])
            documento.add_paragraph(text='<PREENCHA AQUI>', style=self.estilos["texto_informacao"])

        return documento

    def montar_pagina_fase(self, documento: docx.Document, fase: str) -> docx.Document:
        """
        Adiciona uma página de título à caderneta para dividir os semestres do mapeamento geológico.
        :param documento: O documento.
        :param fase: "Mapeamento Geológico I" ou "Mapeamento Geológico II".
        :returns: O documento com a página de título do semestre.
        """
        ic(fase)

        try:  # Quando não há folha de rosto, o documento está inicialmente vazio, e isso causa um IndexError
            documento.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        except IndexError:
            pass
        for i in range(0, 18):
            documento.add_paragraph(text='', style=self.estilos["normal"])
        documento.add_heading(text=fase, level=1)

        return documento

    def montar_pagina_ponto(self, documento: docx.Document, linha: pandas.core.frame.pandas) -> docx.Document:
        """
        Acrescenta uma página de informações de um ponto à caderneta.
        :param documento: O documento
        :param linha: Duplas de rótulos e valores da linha do DataFrame (gerado via DataFrame.itertuples()).
        :returns: O documento com a página do ponto.
        """
        ic(linha.Ponto)

        # Título do ponto
        documento.add_heading(text=linha.Ponto, level=2)

        # Dados que são omitidos em pontos de controle
        und_geo_1 = linha.Unidade_geologica_1 if not pandas.isna(linha.Unidade_geologica_1) else "<Insira aqui a unidade>"
        und_geo_2 = linha.Unidade_geologica_2 if not pandas.isna(linha.Unidade_geologica_2) else None
        und_geo = f"{und_geo_1} / {und_geo_2}" if (und_geo_1 and und_geo_2) and (und_geo_1 != und_geo_2) else und_geo_1

        und_lito_1 = linha.Unidade_litoestratigrafica_1 if not pandas.isna(linha.Unidade_litoestratigrafica_1) else "<Insira aqui a unidade>"
        und_lito_2 = linha.Unidade_litoestratigrafica_2 if not pandas.isna(linha.Unidade_litoestratigrafica_2) else None
        und_lito = f"{und_lito_1} / {und_lito_2}" if (und_lito_1 and und_lito_2) and (und_lito_1 != und_lito_2) else und_lito_1

        tipo_afloramento = linha.Tipo_de_afloramento if not pandas.isna(linha.Tipo_de_afloramento) else "-"
        in_situ = linha.In_situ if not pandas.isna(linha.In_situ) else "-"
        grau_intemperismo = linha.Grau_de_intemperismo if not pandas.isna(linha.Grau_de_intemperismo) else "-"
        amostras = linha.Numero_de_amostras if linha.Numero_de_amostras > 0 else "-"

        # Dicionário com informações que irão para a tabela de cabeçalho
        dados_tabela = {
            'DATA:': f"{linha.Data}",
            'COORDENADAS:': f"{linha.Easting:.0f} E {linha.Northing:.0f} N   {linha.SRC}",
            'ALTITUDE:': f"{linha.Altitude:.0f} m" if not pandas.isna(linha.Altitude) else "-",
            'MUNICÍPIO:': f"{linha.Municipio} - {linha.UF}",
            'TOPONÍMIA:': f"{linha.Toponimia}" if not pandas.isna(linha.Toponimia) else "-",
            'EQUIPE:': f"{linha.Equipe}",
            'PONTO DE CONTROLE:': f"{linha.Ponto_de_controle}",
            'TIPO DE AFLORAMENTO:': f"{tipo_afloramento if linha.Ponto_de_controle == "Não" else "-"}",
            'IN SITU:': f"{in_situ if linha.Ponto_de_controle == "Não" else "-"}",
            'GRAU DE INTEMPERISMO:': f"{grau_intemperismo if linha.Ponto_de_controle == "Não" else "-"}",
            'AMOSTRAS:': f"{amostras if linha.Ponto_de_controle == "Não" else "-"}",
            'UNIDADE GEOLÓGICA:': f"{und_geo if linha.Ponto_de_controle == "Não" else "-"}",
            'UNIDADE LITOESTRATIGRÁFICA:': f"{und_lito if linha.Ponto_de_controle == "Não" else "-"}",
        }

        # Preenche a tabela de cabeçalho
        table = documento.add_table(rows=0, cols=2)
        table.style = self.estilos["tabela_cabecalho"]
        for key in dados_tabela.keys():
            lin = table.add_row().cells
            # Coluna esquerda
            lin[0].text = key
            lin[0].paragraphs[0].style = self.estilos["tabela_esquerda"]
            # Coluna direita
            lin[1].text = dados_tabela[key]
            lin[1].paragraphs[0].style = self.estilos["tabela_direita"]

        # Ajusta a largura das colunas da tabela
        for celula in table.columns[0].cells:
            celula.width = docx.shared.Inches(2.1)
        for celula in table.columns[1].cells:
            celula.width = docx.shared.Inches(3.8)

        # Adiciona a seção de descrição do ponto
        documento.add_paragraph(text="DESCRIÇÃO", style=self.estilos["subtitulo"])
        documento.add_paragraph(text="<Descrição do ponto aqui>", style=self.estilos["normal"])

        # Se for um ponto de controle, encerra aqui
        if linha.Ponto_de_controle == "Sim":
            return documento

        # Adiciona a seção de amostras
        documento.add_paragraph(text="AMOSTRAS", style=self.estilos["subtitulo"])
        if linha.Numero_de_amostras > 0:
            abc = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
            for i in range(0, linha.Numero_de_amostras):
                letra = abc[i]
                documento.add_paragraph(
                    text=f"• {linha.Ponto}{letra}: <Descrição da amostra aqui>",
                    style=self.estilos["normal"]
                )
        else:
            documento.add_paragraph(text=f"• XXX-YYYYZ: <Descrição da amostra aqui>", style=self.estilos["normal"])
            documento.add_paragraph(text=f"REMOVA esta seção caso não haja amostras.", style=self.estilos["anotacao"])

        # Adiciona a seção de medidas estruturais
        documento.add_paragraph(text="MEDIDAS ESTRUTURAIS", style=self.estilos["subtitulo"])
        documento.add_paragraph(text="• <sigla> = <medida>", style=self.estilos["normal"])
        documento.add_paragraph(
            text=f"Use a notação xxx/yy para estruturas planares e yy-xxx para estruturas lineares, onde xxx = sentido "
                 f"de mergulho (dip direction ou trend) e yy = ângulo de mergulho (dip ou plunge). Ex: Lb = 20-180.",
            style=self.estilos["anotacao"]
        )
        documento.add_paragraph(text=f"REMOVA esta seção caso não haja medidas.", style=self.estilos["anotacao"])

        # Adiciona a seção de croquis
        documento.add_paragraph(text="CROQUIS", style=self.estilos["subtitulo"])
        documento.add_paragraph(
            text="<Insira aqui os croquis elaborados para o afloramento e suas respectivas legendas>",
            style=self.estilos["normal"]
        )
        documento.add_paragraph(text=f"REMOVA esta seção caso não haja croquis.", style=self.estilos["anotacao"])

        # Adiciona a seção de fotos
        documento.add_paragraph(text="FOTOS", style=self.estilos["subtitulo"])
        documento.add_paragraph(
            text="<Insira aqui os painéis de fotos tiradas no afloramento e suas respectivas legendas>",
            style=self.estilos["normal"]
        )
        documento.add_paragraph(text=f"REMOVA esta seção caso não haja fotos.", style=self.estilos["anotacao"])

        return documento

    def salvar_caderneta(self, caminho: str):
        """
        Salva a caderneta como um arquivo .docx.
        :param caminho: O caminho do arquivo.
        :returns: Nada.
        """
        ic(caminho)

        self.caderneta.core_properties.title = "Caderneta de Campo Compilada"
        self.caderneta.core_properties.author = "Template Builder"
        self.caderneta.core_properties.category = "Relatório Técnico"
        self.caderneta.core_properties.created = datetime.now()
        self.caderneta.core_properties.keywords = "Geologia, Mapeamento Geológico"
        self.caderneta.core_properties.language = "Português (Brasil)"
        self.caderneta.core_properties.subject = "Geologia"

        if not caminho.endswith(".docx"):
            caminho += ".docx"

        self.caderneta.save(caminho)
