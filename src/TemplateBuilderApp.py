# -*- coding: utf-8 -*-
"""
Created on Thu Aug 18 10:54:53 2022

@author: Gabriel Maccari
"""

import sys
from os import getcwd as os_getcwd
from PyQt6.QtWidgets import (QApplication, QWidget, QPushButton, QLabel,
                             QFileDialog, QFrame, QMessageBox)
from PyQt6.QtGui import QIcon, QFont
import pandas
import docx #python-docx

# Monta a GUI do programa
class MyApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Template Builder')
        self.setMinimumSize(500, 315)
        self.setMaximumSize(500, 315)
        self.setWindowIcon(QIcon('icons/book.ico'))

        self.filePromptLabel = QLabel('Selecione um arquivo .xlsx contendo os '+
                                      'dados dos pontos.', self)
        self.filePromptLabel.setGeometry(5, 5, 415, 25)

        self.fileButton = QPushButton('Selecionar', self)
        self.fileButton.setGeometry(420, 5, 75, 25)
        self.fileButton.clicked.connect(self.open_file)

        line1 = QFrame(self)
        line1.setGeometry(5, 35, 490, 3)
        line1.setLineWidth(1)
        line1.setFrameShape(QFrame.Shape.HLine)
        line1.setFrameShadow(QFrame.Shadow.Sunken)

        self.statusLabels = []

        c1x1, c1x2 = 5, 150
        c2x1, c2x2 = 250, 400
        w1, w2 = 150, 95
        y0 = 40
        y = y0

        lbPonto = QLabel('Ponto', self)
        lbPonto.setGeometry(c1x1, y, w1, 20)
        stPonto = QLabel('-', self)
        stPonto.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stPonto)
        y += 20
        lbEasting = QLabel('Easting', self)
        lbEasting.setGeometry(c1x1, y, w1, 20)
        stEasting = QLabel('-', self)
        stEasting.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stEasting)
        y += 20
        lbNorthing = QLabel('Northing', self)
        lbNorthing.setGeometry(c1x1, y, w1, 20)
        stNorthing = QLabel('-', self)
        stNorthing.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stNorthing)
        y += 20
        lbSRC = QLabel('SRC', self)
        lbSRC.setGeometry(c1x1, y, w1, 20)
        stSRC = QLabel('-', self)
        stSRC.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stSRC)
        y += 20
        lbAltitude = QLabel('Altitude', self)
        lbAltitude.setGeometry(c1x1, y, w1, 20)
        stAltitude = QLabel('-', self)
        stAltitude.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stAltitude)
        y += 20
        lbDisciplina = QLabel('Disciplina', self)
        lbDisciplina.setGeometry(c1x1, y, w1, 20)
        stDisciplina = QLabel('-', self)
        stDisciplina.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stDisciplina)
        y += 20
        lbData = QLabel('Data', self)
        lbData.setGeometry(c1x1, y, w1, 20)
        stData = QLabel('-', self)
        stData.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stData)
        y += 20
        lbEquipe = QLabel('Equipe', self)
        lbEquipe.setGeometry(c1x1, y, w1, 20)
        stEquipe = QLabel('-', self)
        stEquipe.setGeometry(c1x2, y, w2, 20)
        self.statusLabels.append(stEquipe)
        ymax = y
        y = y0
        lbToponimia = QLabel('Toponimia', self)
        lbToponimia.setGeometry(c2x1, y, w1, 20)
        stToponimia = QLabel('-', self)
        stToponimia.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stToponimia)
        y += 20
        lbPonto_de_controle = QLabel('Ponto_de_controle', self)
        lbPonto_de_controle.setGeometry(c2x1, y, w1, 20)
        stPonto_de_controle = QLabel('-', self)
        stPonto_de_controle.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stPonto_de_controle)
        y += 20
        lbNumero_de_amostras = QLabel('Numero_de_amostras', self)
        lbNumero_de_amostras.setGeometry(c2x1, y, w1, 20)
        stNumero_de_amostras = QLabel('-', self)
        stNumero_de_amostras.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stNumero_de_amostras)
        y += 20
        lbTipo_de_afloramento = QLabel('Tipo_de_afloramento', self)
        lbTipo_de_afloramento.setGeometry(c2x1, y, w1, 20)
        stTipo_de_afloramento = QLabel('-', self)
        stTipo_de_afloramento.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stTipo_de_afloramento)
        y += 20
        lbIn_situ = QLabel('In_situ', self)
        lbIn_situ.setGeometry(c2x1, y, w1, 20)
        stIn_situ = QLabel('-', self)
        stIn_situ.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stIn_situ)
        y += 20
        lbGrau_de_intemperismo = QLabel('Grau_de_intemperismo', self)
        lbGrau_de_intemperismo.setGeometry(c2x1, y, w1, 20)
        stGrau_de_intemperismo = QLabel('-', self)
        stGrau_de_intemperismo.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stGrau_de_intemperismo)
        y +=20
        lbUnidade = QLabel('Unidade', self)
        lbUnidade.setGeometry(c2x1, y, w1, 20)
        stUnidade = QLabel('-', self)
        stUnidade.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stUnidade)
        y += 20
        lbUnd_lito = QLabel('Unidade_litoestratigrafica', self)
        lbUnd_lito.setGeometry(c2x1, y, w1, 20)
        stUnd_lito = QLabel('-', self)
        stUnd_lito.setGeometry(c2x2, y, w2, 20)
        self.statusLabels.append(stUnd_lito)

        y = ymax
        y += 25

        lbRows = QLabel('Número de pontos:', self)
        lbRows.setGeometry(5, y, 110, 20)
        self.nRows = QLabel('-', self)
        self.nRows.setGeometry(120, y, 35, 20)

        y += 25
        line2 = QFrame(self)
        line2.setGeometry(5, y, 490, 3)
        line2.setLineWidth(1)
        line2.setFrameShape(QFrame.Shape.HLine)
        line2.setFrameShadow(QFrame.Shadow.Sunken)

        y += 5
        self.readyLabel = QLabel('Aguardando carregamento dos dados.', self)
        self.readyLabel.setGeometry(5, y, 490, 20)

        y += 25
        self.buildButton = QPushButton('Gerar template', self)
        self.buildButton.setGeometry(5, y, 490, 35)
        self.buildButton.setEnabled(False)
        self.buildButton.clicked.connect(self.build_template)

        y += 35
        copyrightLabel = QLabel('© 2022 Gabriel Maccari <gabriel.maccari@' +
                                'hotmail.com>', self)
        copyrightLabel.setGeometry(5, y, 340, 20)
        copyrightLabel.setFont(QFont('Sans Serif', 8))

        self.style_template = self.load_dependency('templates/' +
                                                   'style_template.docx')


    # Carrega o documento do word que contém os estilos de formatação da
    # caderneta
    def load_dependency(self, path):
        try:
            return docx.Document(path)
        except Exception as e:
            erro = str(e)
            msg = QMessageBox(parent=self,
                              text=f'Um dos arquivos necessários para o '+
                              'funcionamento da ferramenta está faltando '+
                              '({path}). Restaure o arquivo em questão e tente'+
                              ' novamente.\n\n{erro}')
            msg.setWindowTitle("Erro ao carregar arquivo")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
            self.filePromptLabel.setText('Arquivo de estilos não encontrado.')
            self.filePromptLabel.setStyleSheet('QLabel {color: red}')
            self.fileButton.setEnabled(False)
            return None


    # Abre uma janelinha de seleção de arquivo para o usuário e carrega as
    # informações da tabela selecionada para um DataFrame
    def open_file(self):
        file = QFileDialog.getOpenFileName(window, caption='Selecione uma '+
                'tabela .xlsx contendo os dados de entrada.',
                directory=os_getcwd(), filter='*.xlsx',
                initialFilter='*.xlsx')
        self.file_path = file[0]

        error_lbl = 'Não foi possível abrir o arquivo.'

        if self.file_path != '':
            try:
                # Salva a primeira aba da tabela em um DataFrame
                self.df = pandas.read_excel(self.file_path, engine='openpyxl')

                # Descarta colunas sem nome
                remove_cols = [col for col in self.df.columns
                               if 'Unnamed' in col]
                self.df.drop(remove_cols, axis='columns', inplace=True)
                # Descarta linhas vazias
                self.df.dropna(how='all', axis='index', inplace=True)
                # Verifica se existem linhas preenchidas no arquivo
                rows = len(self.df.index)
                if rows <= 0:
                    error_lbl = 'O arquivo selecionado está vazio.'
                    raise Exception('File contains zero rows')
                else:
                    file_name = self.file_path.split('/')
                    self.filePromptLabel.setText(file_name[-1])
                    self.filePromptLabel.setStyleSheet('QLabel {color: green}')
                    self.nRows.setText(str(rows))
                    self.check_columns()

            except:
                self.filePromptLabel.setText(error_lbl)
                self.filePromptLabel.setStyleSheet('QLabel {color: red}')
                for lbl in self.statusLabels:
                    lbl.setText('-')
                    lbl.setStyleSheet('QLabel {color: black}')
                self.nRows.setText('-')
                self.readyLabel.setText('Aguardando carregamento dos dados.')
                self.readyLabel.setStyleSheet('QLabel {color: black}')
                self.buildButton.setEnabled(False)

            try:
                self.df['Data'] = self.df['Data'].astype(str)
            except:
                self.df['Data'] = self.df['Data'].dt.strftime('%d/%m/%Y')

        return


    # Verifica se os dados contidos em cada coluna estão de acordo com o padrão
    def check_columns(self):
        essential_columns = {'Ponto':
                                {'name': 'Ponto',
                                 'dtype': 'object',
                                 'null_allowed': False,
                                 'domain':[]},
                             'Easting':
                                 {'name': 'Easting',
                                  'dtype': 'float64',
                                  'null_allowed': False,
                                  'domain':[]},
                             'Northing':
                                 {'name': 'Northing',
                                  'dtype': 'float64',
                                  'null_allowed': False,
                                  'domain':[]},
                             'SRC':
                                 {'name': 'SRC',
                                  'dtype': 'object',
                                  'null_allowed': False,
                                  'domain':[]},
                             'Altitude':
                                 {'name': 'Altitude',
                                  'dtype': 'float64',
                                  'null_allowed': True,
                                  'domain':[]},
                             'Disciplina':
                                 {'name': 'Disciplina',
                                  'dtype': 'object',
                                  'null_allowed': False,
                                  'domain':['Mapeamento Geológico I',
                                            'Mapeamento Geológico II']},
                             'Data':
                                 {'name': 'Data',
                                  'dtype': 'datetime64[ns]',
                                  'null_allowed': False,
                                  'domain':[]},
                             'Equipe':
                                 {'name': 'Equipe',
                                  'dtype': 'object',
                                  'null_allowed': False,
                                  'domain':[]},
                             'Toponimia':
                                 {'name': 'Toponimia',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':[]},
                             'Ponto_de_controle':
                                 {'name': 'Ponto_de_controle',
                                  'dtype': 'object',
                                  'null_allowed': False,
                                  'domain':['Sim','Não']},
                             'Numero_de_amostras':
                                 {'name': 'Numero_de_amostras',
                                  'dtype': 'int64',
                                  'null_allowed': False,
                                  'domain':[]},
                             'Tipo_de_afloramento':
                                 {'name': 'Tipo_de_afloramento',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':[]},
                             'In_situ':
                                 {'name': 'In_situ',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':['Sim','Não']},
                             'Grau_de_intemperismo':
                                 {'name': 'Grau_de_intemperismo',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':['Baixo','Médio','Alto']},
                             'Unidade':
                                 {'name': 'Unidade',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':[]},
                             'Unidade_litoestratigrafica':
                                 {'name': 'Unidade_litoestratigrafica',
                                  'dtype': 'object',
                                  'null_allowed': True,
                                  'domain':[]},
                             }

        columns = self.df.columns.to_list()
        self.columns_ok = True

        i = 0
        for c in essential_columns.keys():
            # Se a coluna essencial c existir na tabela
            if c in columns:
                # Tenta converter a coluna para o tipo de dado indicado
                try:
                    dtype = essential_columns[c]['dtype']
                    self.df[c] = self.df[c].astype(dtype, errors='raise')
                    # Se der certo, muda a label de status para verde
                    self.statusLabels[i].setText('OK')
                    self.statusLabels[i].setStyleSheet('QLabel {color: green}')

                # Caso não dê pra converter a coluna para o tipo de dado
                # desejado, significa que há dados incorretos. Nesse caso, muda
                # a label de status para laranja e bloqueia o avanço
                except ValueError:
                    self.statusLabels[i].setText('Fora de formato')
                    self.statusLabels[i].setStyleSheet('QLabel {color: orange}')
                    self.columns_ok = False

                # Caso a coluna não permita nulos e haja células vazias,
                # muda a label de status para laranja e bloqueia o avanço
                if not essential_columns[c]['null_allowed'] and \
                        self.df[c].isnull().values.any():
                    self.statusLabels[i].setText('Células vazias')
                    self.statusLabels[i].setStyleSheet('QLabel {color: '+
                                                       'orange}')
                    self.columns_ok = False

                try:
                    # Verifica se os campos que possuem entrada limitada estão
                    # preenchidos apenas com os valores permitidos
                    if len(essential_columns[c]['domain']) > 0:
                        column_values = self.df[c]
                        if essential_columns[c]['null_allowed']==True:
                            column_values.dropna(inplace=True)
                        if not column_values.isin(
                                essential_columns[c]['domain']).all():
                            raise ValueError(f'Campo "{c}" deve ser '
                                f'preenchido com um dos seguintes valores: '
                                f'{essential_columns[c]["domain"]}.'
                            )

                except ValueError:
                    self.statusLabels[i].setText('Valores inválidos')
                    self.statusLabels[i].setStyleSheet('QLabel {color: orange}')
                    self.columns_ok = False

            # Se a coluna estiver faltando na tabela, muda a label de status
            # para vermelho e bloqueia o avanço
            else:
                self.statusLabels[i].setText('Coluna faltando')
                self.statusLabels[i].setStyleSheet('QLabel {color: red}')
                self.columns_ok = False
            i += 1

            if self.columns_ok:
                self.readyLabel.setText('Clique no botão abaixo para montar o '+
                                        'template da caderneta com os dados '+
                                        'carregados.')
                self.readyLabel.setStyleSheet('QLabel {color: green}')
                self.buildButton.setEnabled(True)
            else:
                self.readyLabel.setText('Corrija na tabela os problemas '+
                                        'indicados acima e carregue novamente '+
                                        'o arquivo.')
                self.readyLabel.setStyleSheet('QLabel {color: orange}')
                self.buildButton.setEnabled(False)

        return


    # Monta o arquivo do word
    def build_template(self):

        # Recarrega o template, por garantia
        self.style_template = self.load_dependency('templates/' +
                                                   'style_template.docx')

        document = None
        document = self.style_template
        dataframe = self.df

        # Pega os estilos do template
        normal = document.styles['Normal']
        title = document.styles['Title']
        heading1 = document.styles['Heading 1']
        heading2 = document.styles['Heading 2']
        subtitle = document.styles['Subtitle']
        info_title = document.styles['Título de informação']
        info_text = document.styles['Texto de informação']
        caption = document.styles['Caption']
        table_title = document.styles['Tabela - Coluna esquerda']
        table_text = document.styles['Tabela - Coluna direita']
        header_table = document.styles['Tabela de cabeçalho']

        # Cria uma lista com todas as colunas do dataframe
        column_list = dataframe.columns.to_list()

        expected_columns = ['Ponto','Disciplina','SRC','Easting','Northing',
                'Altitude','Toponimia','Data','Equipe','Ponto_de_controle',
                'Numero_de_amostras','Possui_croquis','Possui_fotos',
                'Tipo_de_afloramento','In_situ','Grau_de_intemperismo',
                'Unidade','Unidade_litoestratigrafica']

        field_notebook_table = (True if column_list[0:18] == expected_columns
                                else False)
        if field_notebook_table:
            msr_columns = (column_list[18:] if len(column_list)<=32
                           else column_list[18:33])

        # Retira as timestamps das datas
        dataframe['Data'] = pandas.to_datetime(dataframe.Data,
                                               errors='coerce',
                                               dayfirst=True)
        dataframe['Data'] = dataframe['Data'].dt.strftime('%d/%m/%Y')

        # Deleta o primeiro parágrafo do template
        paragraph = document.paragraphs[0]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

        # Monta a primeira página da caderneta
        for i in range(0, 15):
            if i<=9 or (i>10 and i<13) or i>13:
                document.add_paragraph(text='', style=normal)
            elif i==10:
                document.add_paragraph(text='CADERNETA DE CAMPO COMPILADA',
                                       style=title)
            elif i==13:
                document.add_paragraph(text='MAPEAMENTO GEOLÓGICO UFSC',
                                       style=info_title)

        info_list = ['PROJETO:',
                     'ANO:',
                     'PROFESSORES RESPONSÁVEIS:',
                     'NÚMERO DA ÁREA/FAIXA:',
                     'INTEGRANTES DO GRUPO:'
                    ]

        for info in info_list:
            document.add_paragraph(text=info, style=info_title)
            document.add_paragraph(text='<PREENCHA AQUI>', style=info_text)

        disciplina = ['Mapeamento Geológico I', 'Mapeamento Geológico II']
        d = 0
        ponto = 'undefined'

        # Monta o template
        for point in dataframe.itertuples():
            try:
                ponto = str(point.Ponto)

                # Adiciona uma página de título para o semestre/disciplina
                if d < 2 and point.Disciplina == disciplina[d]:
                    document.paragraphs[-1].add_run().add_break(
                        docx.enum.text.WD_BREAK.PAGE)
                    for i in range(0, 18):
                        document.add_paragraph(text='', style=normal)
                    document.add_heading(text=point.Disciplina, level=1)
                    d += 1

                # Quebra a página antes do título do ponto
                document.paragraphs[-1].add_run().add_break(
                    docx.enum.text.WD_BREAK.PAGE)

                # Título do ponto
                document.add_heading(text=ponto, level=2)

                # Formata os dados que irão para o cabeçalho
                data = (point.Data if str(point.Data) != 'NaT' else '-')
                coordenadas = (('%.0fE %.0fN   %s' %
                                (point.Easting, point.Northing, point.SRC)
                                if (str(point.Easting) != 'nan'
                                and str(point.Northing) != 'nan')
                                else '______E _______ N'))
                altitude = ('%.0f m' % (point.Altitude)
                            if str(point.Altitude) != 'nan'
                            else '-')
                topo = (point.Toponimia if point.Toponimia != 'nan' else '-')
                amostras = (('%d' % (point.Numero_de_amostras))
                            if (str(point.Numero_de_amostras).isnumeric()
                            and point.Numero_de_amostras != 0)
                            else '-')
                und = ((str(point.Unidade) + ' - ' +
                       str(point.Unidade_litoestratigrafica))
                       if (str(point.Unidade) != 'nan'
                       and str(point.Unidade_litoestratigrafica) != 'nan')
                       else '-')

                # Dicionário com informações que irão para a tabela de cabeçalho
                table_data = {
                    'DATA:': data,
                    'COORDENADAS:': coordenadas,
                    'ALTITUDE:': altitude,
                    'TOPONÍMIA:': topo,
                    'EQUIPE:': point.Equipe,
                    'PONTO DE CONTROLE:': point.Ponto_de_controle,
                    'TIPO DE AFLORAMENTO:': point.Tipo_de_afloramento,
                    'IN SITU:': point.In_situ,
                    'GRAU DE INTEMPERISMO:': point.Grau_de_intemperismo,
                    'AMOSTRAS:': amostras,
                    'UNIDADE:': und
                }

                # Monta a tabela de cabeçalho
                table = document.add_table(rows=0, cols=2)
                for key in table_data.keys():
                    row = table.add_row().cells
                    row[0].text = key
                    row[0].paragraphs[0].style = table_title
                    row[1].text = str(table_data[key]) \
                        if str(table_data[key]) != 'nan' \
                        else '-'
                    row[1].paragraphs[0].style = table_text
                table.style = header_table
                for cell in table.columns[0].cells:
                    cell.width = docx.shared.Inches(2.1)
                for cell in table.columns[1].cells:
                    cell.width = docx.shared.Inches(3.8)

                # Adiciona a seção de descrição do ponto
                document.add_paragraph(text='DESCRIÇÃO', style=subtitle)
                document.add_paragraph(text="...", style=normal)

                if point.Ponto_de_controle != 'Sim' \
                        and point.Ponto_de_controle != 'sim':
                    # Adiciona a seção de amostras, se houver alguma
                    try:
                        sample_qty = int(point.Numero_de_amostras)
                        if sample_qty > 0:
                            document.add_paragraph(text='AMOSTRAS',
                                    style=subtitle)
                            abc = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
                            for i in range(0, sample_qty):
                                document.add_paragraph(text=('• %s%c: ...' \
                                        % (point.Ponto, abc[i])), style=normal)
                    except:
                        pass

                    # Caso a tabela carregada seja a tabela original da
                    # caderneta, procura os campos de estruturas para preencher
                    if field_notebook_table:
                        measurements = []
                        # Itera nos campos de medidas
                        for i, column_name in enumerate(msr_columns):
                            # Conteúdo do campo
                            msr = str(point[i+19])
                            # Se não for uma célula vazia
                            if msr != '' and msr != 'nan':
                                # Procura uma sigla entre parênteses
                                if '(' in column_name and ')' in column_name:
                                    initials = column_name[column_name.find("(")
                                               + 1:column_name.find(")")]
                                # Se não encontrar sigla, usa o nome da coluna
                                else:
                                    initials = column_name.replace('_', ' ')
                                # Adiciona as medidas a uma lista
                                measurements.append(f'• {initials} =' +
                                                    f' {msr}')
                        if len(measurements)>0:
                            # Adiciona a seção de medidas
                            document.add_paragraph(text='MEDIDAS ESTRUTURAIS',
                                                   style=subtitle)
                            for m in measurements:
                                document.add_paragraph(text=m, style=normal)

                    # Se for uma tabela diferente (provavelmente a do BD),
                    # preenche com texto padrão
                    else:
                        # Adiciona a seção de medidas
                        document.add_paragraph(text='MEDIDAS ESTRUTURAIS',
                                               style=subtitle)
                        document.add_paragraph(text='• S0 = 000/00',
                                               style=normal)
                        document.add_paragraph(text='• Pc = 00-000',
                                               style=normal)
                        document.add_paragraph(
                                text='<Preencha aqui as medidas '
                                'estruturais coletadas no afloramento. Remova '+
                                'esta seção caso não haja medidas estruturais.'+
                                ' Use sentido/mergulho (Ex: 220/30) para '+
                                'medidas de estruturas planares e mergulho-'+
                                'sentido (Ex: 20-340) para medidas de '+
                                'estruturas lineares. Acrescente um "?" ao '+
                                'final de medidas duvidosas.>',
                                style=normal)


                    # Caso o campo Possui_croquis exista na tabela, utiliza ele
                    # para determinar se coloca ou não a seção de croquis
                    if 'Possui_croquis' in column_list:
                        has_sketches = (False if point.Possui_croquis == 'Não'
                                              or point.Possui_croquis == 'não'
                                              or point.Possui_croquis == 'nao'
                                        else True)
                    else:
                        has_sketches = True

                    # Adiciona a seção de croquis
                    if has_sketches:
                        document.add_paragraph(text='CROQUIS', style=subtitle)
                        document.add_paragraph(
                                text='<Insira aqui os croquis '+
                                'elaborados para o afloramento. Remova esta '+
                                'seção caso não haja croquis.>',
                                style=normal)

                    # Mesma coisa para as fotos
                    if 'Possui_fotos' in column_list:
                        has_photos = (False if point.Possui_fotos == 'Não'
                                            or point.Possui_fotos == 'não'
                                            or point.Possui_fotos == 'nao'
                                      else True)
                    else:
                        has_photos = True

                    # Adiciona a seção de fotos
                    if has_photos:
                        document.add_paragraph(text='FOTOS', style=subtitle)
                        document.add_paragraph(
                                text='<Insira aqui os painéis '+
                                'de fotos tiradas no afloramento. '+
                                'Remova esta seção caso não haja fotos.>',
                                style=normal)

            # Caso haja algum problema ao montar o template, tenta identificar
            # em qual ponto está o problema e mostra ao usuário
            except Exception as e:
                erro = str(e)
                if ponto != 'undefined':
                    msg = QMessageBox(
                            parent=self, text='Ocorreu um erro ao '+
                            'montar o template. Verifique e corrija os dados '+
                            f'do ponto {ponto} e tente novamente.\n\n{erro}')
                else:
                    msg = QMessageBox(
                            parent=self, text=f'Ocorreu um erro ao '+
                            'montar o template. Verifique os dados e tente '+
                            f'novamente.\n\n{erro}')
                msg.setWindowTitle("Erro ao preencher o template")
                msg.setIcon(QMessageBox.Icon.Critical)
                msg.exec()
                return

        self.save_file(document)

        return


    #Abre uma janelinha para o usuário salvar o arquivo do word
    def save_file(self, document):
        # Tenta salvar o arquivo
        self.outFile = QFileDialog.getSaveFileName(self, 'Salvar template',
                str(os_getcwd()), 'Documento do Word (*.docx)')
        output_path = self.outFile[0]

        if output_path != '':
            try:
                document.save(output_path)
                msg = QMessageBox(parent=self, text='Template criado com '+
                        f'sucesso!')
                msg.setWindowTitle("Sucesso")
            except:
                msg = QMessageBox(parent=self, text='Ocorreu um erro ao salvar'+
                        f' o arquivo. Verifique se "{output_path}" não está '+
                        'aberto em outro programa e se você possui permissão '+
                        'para salvar nessa pasta, e então tente novamente.')
                msg.setWindowTitle("Erro ao salvar o arquivo")
                msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()


# Main
if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MyApp()
    window.show()
    sys.exit(app.exec())
