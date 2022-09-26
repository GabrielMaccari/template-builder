# -*- coding: utf-8 -*-
"""
Created on Thu Aug 18 10:54:53 2022

@author: Gabriel Maccari
"""

import sys
from os import getcwd as os_getcwd
from PyQt6.QtWidgets import QApplication, QWidget, QPushButton, QLabel, QFileDialog, QFrame, QMessageBox
from PyQt6.QtGui import QIcon, QFont
import pandas
import docx

class MyApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Template Builder')
        self.setMinimumSize(500, 335)
        self.setMaximumSize(500, 335)
        self.setWindowIcon(QIcon('icons/book.ico'))
        
        self.filePromptLabel = QLabel('Selecione um arquivo .xlsx contendo os dados dos pontos.', self)
        self.filePromptLabel.setGeometry(5, 5, 415, 25)

        self.fileButton = QPushButton('Selecionar', self)
        self.fileButton.setGeometry(420, 5, 75, 25)
        self.fileButton.clicked.connect(self.open_file)
        
        self.line1 = QFrame(self)
        self.line1.setGeometry(5, 35, 490, 3)
        self.line1.setLineWidth(1)
        self.line1.setFrameShape(QFrame.Shape.HLine)
        self.line1.setFrameShadow(QFrame.Shadow.Sunken)

        c1x1, c1x2 = 5, 150
        c2x1, c2x2 = 250, 400
        w1, w2 = 150, 95
        y0 = 40
        y = y0
        
        self.lbPonto = QLabel('Ponto', self)
        self.lbPonto.setGeometry(c1x1, y, w1, 20)
        self.stPonto = QLabel('-', self)
        self.stPonto.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbUTM_E = QLabel('UTM_E', self)
        self.lbUTM_E.setGeometry(c1x1, y, w1, 20)
        self.stUTM_E = QLabel('-', self)
        self.stUTM_E.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbUTM_N = QLabel('UTM_N', self)
        self.lbUTM_N.setGeometry(c1x1, y, w1, 20)
        self.stUTM_N = QLabel('-', self)
        self.stUTM_N.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbAltitude = QLabel('Altitude', self)
        self.lbAltitude.setGeometry(c1x1, y, w1, 20)
        self.stAltitude = QLabel('-', self)
        self.stAltitude.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbDisciplina = QLabel('Disciplina', self)
        self.lbDisciplina.setGeometry(c1x1, y, w1, 20)
        self.stDisciplina = QLabel('-', self)
        self.stDisciplina.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbData = QLabel('Data', self)
        self.lbData.setGeometry(c1x1, y, w1, 20)
        self.stData = QLabel('-', self)
        self.stData.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbEquipe = QLabel('Equipe', self)
        self.lbEquipe.setGeometry(c1x1, y, w1, 20)
        self.stEquipe = QLabel('-', self)
        self.stEquipe.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbToponimia = QLabel('Toponimia', self)
        self.lbToponimia.setGeometry(c1x1, y, w1, 20)
        self.stToponimia = QLabel('-', self)
        self.stToponimia.setGeometry(c1x2, y, w2, 20)
        y+=20
        self.lbPonto_de_controle = QLabel('Ponto_de_controle', self)
        self.lbPonto_de_controle.setGeometry(c1x1, y, w1, 20)
        self.stPonto_de_controle = QLabel('-', self)
        self.stPonto_de_controle.setGeometry(c1x2, y, w2, 20)
        ymax=y
        y=y0
        self.lbUnidade = QLabel('Unidade', self)
        self.lbUnidade.setGeometry(c2x1, y, w1, 20)
        self.stUnidade = QLabel('-', self)
        self.stUnidade.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbUnidade_litoestratigrafica = QLabel('Unidade_litoestratigrafica', self)
        self.lbUnidade_litoestratigrafica.setGeometry(c2x1, y, w1, 20)
        self.stUnidade_litoestratigrafica = QLabel('-', self)
        self.stUnidade_litoestratigrafica.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbTipo_de_afloramento = QLabel('Tipo_de_afloramento', self)
        self.lbTipo_de_afloramento.setGeometry(c2x1, y, w1, 20)
        self.stTipo_de_afloramento = QLabel('-', self)
        self.stTipo_de_afloramento.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbIn_situ = QLabel('In_situ', self)
        self.lbIn_situ.setGeometry(c2x1, y, w1, 20)
        self.stIn_situ = QLabel('-', self)
        self.stIn_situ.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbGrau_de_intemperismo = QLabel('Grau_de_intemperismo', self)
        self.lbGrau_de_intemperismo.setGeometry(c2x1, y, w1, 20)
        self.stGrau_de_intemperismo = QLabel('-', self)
        self.stGrau_de_intemperismo.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbNumero_de_amostras = QLabel('Numero_de_amostras', self)
        self.lbNumero_de_amostras.setGeometry(c2x1, y, w1, 20)
        self.stNumero_de_amostras = QLabel('-', self)
        self.stNumero_de_amostras.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbPossui_croquis = QLabel('Possui_croquis', self)
        self.lbPossui_croquis.setGeometry(c2x1, y, w1, 20)
        self.stPossui_croquis = QLabel('-', self)
        self.stPossui_croquis.setGeometry(c2x2, y, w2, 20)
        y+=20
        self.lbPossui_fotos = QLabel('Possui_fotos', self)
        self.lbPossui_fotos.setGeometry(c2x1, y, w1, 20)
        self.stPossui_fotos = QLabel('-', self)
        self.stPossui_fotos.setGeometry(c2x2, y, w2, 20)
        
        self.labels = [self.stPonto, self. stUTM_E, self.stUTM_N, self.stAltitude, self.stDisciplina, self.stData, self.stEquipe, self.stToponimia, self.stPonto_de_controle, self.stUnidade, self.stUnidade_litoestratigrafica, self.stTipo_de_afloramento, self.stIn_situ, self.stGrau_de_intemperismo, self.stNumero_de_amostras, self.stPossui_croquis, self.stPossui_fotos]
        
        y=ymax
        y+=25
        self.lbOtherColumns = QLabel('Outras colunas:', self)
        self.lbOtherColumns.setGeometry(5, y, 95, 20)
        self.nOtherColumns = QLabel('-', self)
        self.nOtherColumns.setGeometry(95, y, 35, 20)
        
        self.lbRows = QLabel('Número de pontos:', self)
        self.lbRows.setGeometry(140, y, 105, 20)
        self.nRows = QLabel('-', self)
        self.nRows.setGeometry(250, y, 35, 20)
        
        y+=25
        self.line2 = QFrame(self)
        self.line2.setGeometry(5, y, 490, 3)
        self.line2.setLineWidth(1)
        self.line2.setFrameShape(QFrame.Shape.HLine)
        self.line2.setFrameShadow(QFrame.Shadow.Sunken)
        
        y+=5
        self.readyLabel = QLabel('Aguardando carregamento da planilha de dados.', self)
        self.readyLabel.setGeometry(5, y, 490, 20)
        
        y+=25
        self.buildButton = QPushButton('Gerar template', self)
        self.buildButton.setGeometry(5, y, 490, 35)
        self.buildButton.setEnabled(False)
        self.buildButton.clicked.connect(self.build_template)
        
        y+=35
        self.copyrightLabel = QLabel('© 2022 Gabriel Maccari <gabriel.maccari@hotmail.com>', self)
        self.copyrightLabel.setGeometry(5, y, 340, 20)
        self.copyrightLabel.setFont(QFont('Sans Serif', 8))
        
        self.style_template = self.load_dependency('templates/style_template.docx', 'docx')
    
    def load_dependency(self, path, file_type):
        try:
            if file_type == 'xlsx':
                return pandas.read_excel(path, engine='openpyxl')
            elif file_type == 'docx':
                return docx.Document(path)
        except Exception as e:
            msg = QMessageBox(parent=self, text=('Um dos arquivos necessários para o funcionamento da ferramenta está faltando (%s). Restaure o arquivo em questão e tente novamente.\n\n%s' % (path, str(e))))
            msg.setWindowTitle("Erro ao carregar arquivo")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
            self.filePromptLabel.setText('Arquivo de estilos não encontrado.')
            self.filePromptLabel.setStyleSheet('QLabel {color: red}')
            self.fileButton.setEnabled(False)
            return None
    
    def open_file(self):
        file = QFileDialog.getOpenFileName(window, caption='Selecione uma tabela .xlsx contendo os dados de entrada.', directory=os_getcwd(), filter='*.xlsx', initialFilter='*.xlsx')
        self.file_path = file[0]
        
        error_lbl = 'Não foi possível abrir o arquivo.'
        
        if self.file_path!='':
            try:
                self.df = pandas.read_excel(self.file_path, engine='openpyxl')
                #Descarta colunas sem nome
                remove_cols = [col for col in self.df.columns if 'Unnamed' in col]
                self.df.drop(remove_cols, axis='columns', inplace=True)  #self.df = self.df.loc[:, ~self.df.columns.str.contains('^Unnamed')]
                #Descarta linhas vazias
                self.df.dropna(how='all', axis='index', inplace=True)
                #Verifica se existem linhas preenchidas no arquivo
                rows = len(self.df.index)
                if rows<=0:
                    error_lbl = 'O arquivo selecionado está vazio.'
                    raise Exception('File contains zero rows')
                else:
                    self.filePromptLabel.setText('Arquivo carregado com sucesso.')
                    self.filePromptLabel.setStyleSheet('QLabel {color: green}')
                    self.nRows.setText(str(rows))
                    self.check_columns()
            except:
                self.filePromptLabel.setText(error_lbl)
                self.filePromptLabel.setStyleSheet('QLabel {color: red}')
                for lbl in self.labels:
                    lbl.setText('-')
                    lbl.setStyleSheet('QLabel {color: black}')
                self.nOtherColumns.setText('-')
                self.nRows.setText('-')
                self.readyLabel.setText('Aguardando carregamento da planilha de dados.')
                self.readyLabel.setStyleSheet('QLabel {color: black}')
                self.buildButton.setEnabled(False)
    
            try:
                self.df['Data'] = self.df['Data'].astype(str)
            except:
                self.df['Data'] = self.df['Data'].dt.strftime('%d/%m/%Y')
            
        return
    
    def check_columns(self):
        essential_columns = {'Ponto':{'name':'Ponto','dtype':'object','null_allowed':False},
                              'UTM_E':{'name':'UTM_E','dtype':'float64','null_allowed':False},
                              'UTM_N':{'name':'UTM_N','dtype':'float64','null_allowed':False},
                              'Altitude':{'name':'Altitude','dtype':'float64','null_allowed':True},
                              'Disciplina':{'name':'Disciplina','dtype':'object','null_allowed':False},
                              'Data':{'name':'Data','dtype':'datetime64','null_allowed':False},
                              'Equipe':{'name':'Equipe','dtype':'object','null_allowed':False},
                              'Toponimia':{'name':'Toponimia','dtype':'object','null_allowed':True},
                              'Ponto_de_controle':{'name':'Ponto_de_controle','dtype':'object','null_allowed':False},
                              'Unidade':{'name':'Unidade','dtype':'object','null_allowed':True},
                              'Unidade_litoestratigrafica':{'name':'Unidade_litoestratigrafica','dtype':'object','null_allowed':True},
                              'Tipo_de_afloramento':{'name':'Tipo_de_afloramento','dtype':'object','null_allowed':True},
                              'In_situ':{'name':'In_situ','dtype':'object','null_allowed':True},
                              'Grau_de_intemperismo':{'name':'Grau_de_intemperismo','dtype':'object','null_allowed':True},
                              'Numero_de_amostras':{'name':'Numero_de_amostras','dtype':'int64','null_allowed':False},
                              'Possui_croquis':{'name':'Possui_croquis','dtype':'object','null_allowed':False},
                              'Possui_fotos':{'name':'Possui_fotos','dtype':'object','null_allowed':False}}
        e_columns = essential_columns.keys()
        
        columns = self.df.columns.to_list()
        self.columns_ok = True
        
        i=0
        for c in e_columns:
            #Se a coluna essencial c existir na tabela
            if c in columns:
                #Tenta converter a coluna para o tipo de dado indicado
                try:
                    dtype = essential_columns[c]['dtype']
                    self.df[c] = self.df[c].astype(dtype, errors='raise')
                    #Se der certo, muda a label de status para verde
                    self.labels[i].setText('OK')
                    self.labels[i].setStyleSheet('QLabel {color: green}')
                    #Caso a coluna não permita nulos e haja células vazias, muda a label de status para laranja e bloqueia o avanço
                    if not essential_columns[c]['null_allowed'] and self.df[c].isnull().values.any():
                        self.labels[i].setText('Há células vazias')
                        self.labels[i].setStyleSheet('QLabel {color: orange}')
                        self.columns_ok = False
                #Caso não dê pra converter a coluna para o tipo de dado desejado, significa que há dados incorretos
                #Nesse caso, muda a label de status para laranja e bloqueia o avanço
                except:
                    self.labels[i].setText('Fora de formato')
                    self.labels[i].setStyleSheet('QLabel {color: orange}')
                    self.columns_ok = False
            #Se a coluna estiver faltando na tabela, muda a label de status para vermelho e bloquea o avanço
            else:
                self.labels[i].setText('Coluna faltando')
                self.labels[i].setStyleSheet('QLabel {color: red}')
                self.columns_ok = False
            i+=1
            
            if self.columns_ok:
                self.readyLabel.setText('Clique no botão abaixo para montar o template da caderneta com os dados carregados.')
                self.readyLabel.setStyleSheet('QLabel {color: green}')
                self.buildButton.setEnabled(True)
            else:
                self.readyLabel.setText('Corrija na tabela os problemas indicados acima e carregue novamente o arquivo.')
                self.readyLabel.setStyleSheet('QLabel {color: orange}')
                self.buildButton.setEnabled(False)
                
        extra_columns = 0
        for c in columns:
            if c not in e_columns:
                extra_columns+=1
        self.nOtherColumns.setText(str(extra_columns))
        
        return
    
    def build_template(self):
        document = self.style_template
        dataframe = self.df
        
        #Pega os estilos do template
        normal = document.styles['Normal']
        title = document.styles['Title']
        heading1 = document.styles['Heading 1']
        heading2 = document.styles['Heading 2']
        subtitle = document.styles['Subtitle']
        info_title = document.styles['Título de informação']
        info_text = document.styles['Texto de informação']
        caption = document.styles['Caption']
        table_title = document.styles['Tabela - Coluna esquerda']
        table_text = document.styles['Tabela - Coluna direita']
        header_table = document.styles['Tabela de cabeçalho']
        
        #Cria uma lista com todas as colunas do dataframe
        column_list = dataframe.columns.to_list() 
        
        #Retira as timestamps das datas
        dataframe['Data'] = pandas.to_datetime(dataframe.Data, errors='coerce')
        dataframe['Data'] = dataframe['Data'].dt.strftime('%d/%m/%Y')
        
        #dataframe['Numero_de_amostras'] = dataframe['Numero_de_amostras'].astype('int64', errors='ignore')

        #Deleta o primeiro parágrafo do template
        paragraph = document.paragraphs[0]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
        
        #Monta a primeira página da caderneta
        for i in range(0,9):
            document.add_paragraph(text='', style=normal)
        document.add_paragraph(text='Caderneta de Campo Compilada', style=title)
        for i in range(0,2):
            document.add_paragraph(text='', style=normal)
        document.add_paragraph(text='Mapeamento Geológico UFSC', style=info_title)
        for i in range(0,2):
            document.add_paragraph(text='', style=normal)
        info_list = ['Projeto:', '<Nome do projeto>', 'Ano:', '<Ano>', 'Professores responsáveis:', '<Professor 1, Professor 2...>', 'Número da área/faixa:', '<Número da faixa>', 'Integrantes do grupo:', '<Integrante 1, Integrante 2...>']
        for i in info_list:
            stl = info_title if i[0]!='<' else info_text
            document.add_paragraph(text=i, style=stl)
        
        disciplina = ['Mapeamento Geológico I','Mapeamento Geológico II']
        d = 0
        ponto = 'undefined'
        
        #Monta o template
        for point in dataframe.itertuples():
            try:
                ponto = str(point.Ponto)
                
                #Adiciona um título para o semestre/disciplina
                if d<2 and point.Disciplina == disciplina[d]:
                    document.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                    for i in range(0,18):
                        document.add_paragraph(text='', style=normal)
                    document.add_heading(text=point.Disciplina, level=1)
                    d+=1
                
                #Quebra a página antes do título do ponto
                document.paragraphs[-1].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                
                #Título do ponto
                document.add_heading(text=ponto, level=2)
                
                #Dicionário com informações que irão para a tabela de cabeçalho
                table_data = {'Data:':point.Data if str(point.Data)!='NaT' else 'nan',
                              'Coordenadas:':('%.0f E / %.0f N' % (point.UTM_E, point.UTM_N) if (str(point.UTM_E)!='nan' and str(point.UTM_N)!='nan') else '______ E / _______ N'),
                              'Altitude:':('%.0f m' % (point.Altitude) if str(point.Altitude)!='nan' else 'nan'),
                              'Equipe:':point.Equipe,
                              'Ponto de controle:':point.Ponto_de_controle,
                              'Tipo de afloramento:':point.Tipo_de_afloramento,
                              'In situ:':point.In_situ,
                              'Grau de intemperismo:':point.Grau_de_intemperismo,
                              'Amostras:':('%d' % (point.Numero_de_amostras)) if str(point.Numero_de_amostras).isnumeric() else 'nan',
                              'Unidade:':(str(point.Unidade)+' - '+str(point.Unidade_litoestratigrafica)) if (str(point.Unidade)!='nan' and str(point.Unidade_litoestratigrafica)!='nan') else 'nan'
                             }
                table_data_keys = table_data.keys()
                
                #Monta a tabela de cabeçalho
                table = document.add_table(rows=0, cols=2)
                for key in table_data_keys:
                    row = table.add_row().cells
                    row[0].text = key
                    row[0].paragraphs[0].style = table_title
                    row[1].text = str(table_data[key]) if str(table_data[key]) != 'nan' else '-'
                    row[1].paragraphs[0].style = table_text
                table.style = header_table
                for cell in table.columns[0].cells:
                    cell.width = docx.shared.Inches(2.1)
                for cell in table.columns[1].cells:
                    cell.width = docx.shared.Inches(3.8)
                
                #Adiciona a seção de toponímia
                document.add_paragraph(text='Toponímia', style=subtitle)
                document.add_paragraph(text=str(point.Toponimia if str(point.Toponimia) != 'nan' else '-'), style=normal)
                
                #Adiciona a seção de descrição do ponto
                document.add_paragraph(text='Descrição', style=subtitle)
                document.add_paragraph(text="...",style=normal)
                
                #Adiciona a seção de amostras, se houver alguma
                try:
                    sample_qty = int(point.Numero_de_amostras)
                    if sample_qty > 0:
                        document.add_paragraph(text='Amostras', style=subtitle)
                        abc = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
                        for i in range(0, sample_qty):
                            document.add_paragraph(text=('• %s%c: ...' % (point.Ponto, abc[i])), style=normal)
                except:
                    pass
                
                #Adiciona a seção de medidas, se houver alguma
                initials_list = []
                msr_list = []
                hasMeasurements = False
                for i in range(18, len(column_list)+1): #Itera em todas as possíveis colunas de medidas (18-x*)   *x ≤ 31 se ninguém fez merda
                    if str(point[i]) != '' and str(point[i]) != 'nan':
                        hasMeasurements = True
                        column = str(column_list[i-1]) #i-1 porque na lista começa em zero, enquanto que no dataframe começa em 1 (o zero é o id)
                        
                        #Adiciona a sigla que está entre parênteses no nome do campo a uma lista de siglas
                        if '(' in column and ')' in column:
                            initials = column[column.find("(")+1:column.find(")")]
                        #Se não tiver nada entre parênteses, adiciona o nome inteiro da coluna
                        else:
                            initials = column.replace('_',' ')
                        initials_list.append(initials)
                        
                        #Adiciona a medida (conteúdo da célula) a uma lista de medidas
                        msr = str(point[i])
                        msr_list.append(msr)
                #Se houver medidas, adiciona elas ao template
                if hasMeasurements:
                    document.add_paragraph(text='Medidas Estruturais', style=subtitle)
                    for x in range(len(msr_list)):
                        document.add_paragraph(text=('• %s = %s' % (initials_list[x], msr_list[x])), style=normal)
                
                #Adiciona a seção de croquis, se houver algum
                if str(point.Possui_croquis) != 'Não' and str(point.Possui_croquis) != '0':
                    document.add_paragraph(text='Croquis', style=subtitle)
                    document.add_paragraph(text='...',style=normal)
                
                #Adiciona a seção de fotos, se houver alguma
                if str(point.Possui_fotos) != 'Não' and str(point.Possui_fotos) != '0':
                    document.add_paragraph(text='Fotos', style=subtitle)
                    document.add_paragraph(text='...',style=normal)
            
            #Caso haja algum problema ao montar o template, tenta identificar em qual ponto está o problema
            except Exception as e:
                if ponto != 'undefined':
                    msg = QMessageBox(parent=self, text=('Ocorreu um erro ao montar o template. Verifique e corrija os dados do ponto %s e tente novamente.\n\n%s' % (ponto, str(e))))
                    msg.setWindowTitle("Erro ao preencher o template")
                    msg.setIcon(QMessageBox.Icon.Critical)
                    msg.exec()
                else:
                    msg = QMessageBox(parent=self, text='Ocorreu um erro ao montar o template. Verifique os dados e tente novamente.\n\n%s' % (str(e)))
                    msg.setWindowTitle("Erro ao preencher o template")
                    msg.setIcon(QMessageBox.Icon.Critical)
                    msg.exec()
                return
            
        self.save_file(document)
            
        return
    
    def save_file(self, document):
        #Tenta salvar o arquivo
        self.outFile = QFileDialog.getSaveFileName(self, 'Salvar template', str(os_getcwd()), '*.docx')
        output_path = self.outFile[0]
        
        if output_path!='':
            try:
                document.save(output_path)
                msg = QMessageBox(parent=self, text=('Template criado com sucesso! O arquivo foi salvo em "%s".' % (output_path)))
                msg.setWindowTitle("Sucesso")
                msg.exec()
            except:
                msg = QMessageBox(parent=self, text=('Ocorreu um erro ao salvar o arquivo. Verifique se "%s" não está aberto em outro programa e se você possui permissão para salvar nessa pasta, e então tente novamente.' % (output_path)))
                msg.setWindowTitle("Erro ao salvar o arquivo")
                msg.setIcon(QMessageBox.Icon.Critical)
                msg.exec()
        
if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MyApp()
    window.show()
    sys.exit(app.exec())